import sys
import os
import re
from collections import defaultdict
from docx import Document
from faker import Faker
import spacy

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

regex_patterns = [
    {'type': 'EMAIL',       'pattern': re.compile(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b')},
    {'type': 'PHONE',       'pattern': re.compile(r'\+?\s*(?:\d{1,3}[\s-]?)?\(?\d{2,4}\)?[-.\ s]?\d{3,4}[-.\ s]?\d{4}\b')},
    {'type': 'SSN',         'pattern': re.compile(r'\b\d{3}-\d{2}-\d{4}\b')},
    {'type': 'CREDIT_CARD', 'pattern': re.compile(r'\b(?:\d[ -]*?){13,16}\b')},
    {'type': 'DATE_OF_BIRTH','pattern': re.compile(r'\b(?:0[1-9]|[12]\d|3[01])[-/.](?:0[1-9]|1[0-2])[-/.](?:19|20)\d{2}\b')},
    {'type': 'IP_ADDRESS',  'pattern': re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')},
    {'type': 'PAN',         'pattern': re.compile(r'\b[A-Z]{5}\d{4}[A-Z]{1}\b')},
    {'type': 'AADHAAR',     'pattern': re.compile(r'\b\d{4}\s?\d{4}\s?\d{4}\b')},
    {'type': 'ADDRESS',     'pattern': re.compile(r'\b(?:Gat|Plot|Tower|Building|Flat|Door|Survey)\s+No\.?\s*[\w/\s,-]+', re.IGNORECASE)},
]

exclude_terms = {
    "company", "prospectus", "schedule", "table", "board", "directors",
    "india", "government", "annexure", "section", "chapter", "act",
    "rupees", "crores", "lakhs", "sebi", "rbi", "limited", "report",
    "offer", "shares", "issue", "public", "bse", "nse", "stock", "exchange",
    "email", "telephone", "tel", "contact", "person", "website", "fax",
    "date", "notice", "time", "period", "year", "fiscal", "total", "amount",
    "price", "bank", "registrar", "manager", "lead", "bidders", "investors",
    "application", "form", "bidding", "bids", "allotment", "equity", "capital",
    "maharashtra", "pune", "mumbai", "delhi", "bengaluru", "chennai", "regulator"
}

def evaluate(input_path):
    print(f"\nLoading document: {input_path}\n")
    nlp = spacy.load("en_core_web_sm")
    doc = Document(input_path)

    hits = defaultdict(list)          
    ner_names = set()
    ner_orgs  = set()
    contact_names = set()

    all_texts = []

    for p in doc.paragraphs:
        if p.text.strip():
            all_texts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_texts.append(cell.text)

    print(f"Total text blocks scanned: {len(all_texts)}\n")

    for text in all_texts:
       
        for item in regex_patterns:
            for m in item['pattern'].finditer(text):
                hits[item['type']].append(m.group(0).strip())

        contact_match = re.search(r'(?:Contact\s*Person[:]?)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', text)
        if contact_match:
            name = contact_match.group(1).strip()
            if name.lower() not in exclude_terms:
                contact_names.add(name)

        spacydoc = nlp(text)
        for ent in spacydoc.ents:
            val = ent.text.strip()
            low = val.lower()
            if low in exclude_terms or val.isdigit() or len(val) <= 3:
                continue
            if val.isupper() and len(val.split()) == 1:
                continue
            if ent.label_ == "PERSON" and len(val.split()) >= 2:
                ner_names.add(val)
            elif ent.label_ == "ORG" and len(val.split()) >= 2:
                if not any(t in low for t in ["sebi", "bse", "nse", "act", "supreme court"]):
                    ner_orgs.add(val)

    hits['PERSON (NER)'] = list(ner_names | contact_names)
    hits['ORG (NER)']    = list(ner_orgs)

   
    print("=" * 65)
    print(f"  {'PII TYPE':<22} {'UNIQUE HITS':>12}  SAMPLE VALUES")
    print("=" * 65)

    total_unique = 0
    total_instances = 0

    type_order = ['EMAIL','PHONE','PAN','AADHAAR','SSN','CREDIT_CARD',
                  'DATE_OF_BIRTH','IP_ADDRESS','ADDRESS','PERSON (NER)','ORG (NER)']

    for t in type_order:
        vals = hits.get(t, [])
        unique_vals = list(dict.fromkeys(vals))   
        sample = ", ".join(f'"{v[:40]}"' for v in unique_vals[:3])
        if len(unique_vals) > 3:
            sample += f" ... +{len(unique_vals)-3} more"
        print(f"  {t:<22} {len(unique_vals):>12}  {sample}")
        total_unique    += len(unique_vals)
        total_instances += len(vals)

    print("=" * 65)
    print(f"  {'TOTAL':<22} {total_unique:>12}  unique PII values detected")
    print(f"  {'TOTAL (with repeats)':<22} {total_instances:>12}  raw regex/NER hits")
    print("=" * 65)

    print("\nKNOWN MISSES (not caught by tool):")
    print("  - PAN card image (embedded in doc)  -> not readable by python-docx")
    print("  - Aadhaar card image (embedded)      -> not readable by python-docx")
    print("  - Single-word names (by design)      -> too many false positives")
    print()


if __name__ == "__main__":
    path = sys.argv[1] if len(sys.argv) > 1 else "input_docs/Red Herring Prospectus (2).docx"
    evaluate(path)
