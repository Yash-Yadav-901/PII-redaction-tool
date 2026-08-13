import os
import sys
import re
from docx import Document
from faker import Faker
import spacy

class PIIDetectionAndRedaction:
    def __init__(self):
        self.fake = Faker()
        self.pii_mapping = {}
        self.discovered_names = set()
        self.discovered_orgs = set()
        
        print("Initializing NLP engine....")
        self.nlp = spacy.load("en_core_web_sm")
        
        # these are terms that are commonly found in financial documents and should not be considered PII, so to exclude them from redaction
        self.exclude_terms = {
            "company", "prospectus", "schedule", "table", "board", "directors", 
            "india", "government", "annexure", "section", "chapter", "act", 
            "rupees", "crores", "lakhs", "sebi", "rbi", "limited", "report",
            "offer", "shares", "issue", "public", "bse", "nse", "stock", "exchange",
            "email", "telephone", "tel", "contact", "person", "website", "fax",
            "date", "notice", "time", "period", "year", "fiscal", "total", "amount",
            "price", "bank", "registrar", "manager", "lead", "bidders", "investors",
            "application", "form", "bidding", "bids", "allotment", "equity", "capital",
            "regulator"
        }
        

        self.regex_patterns = [
            #Email Addresses
            {'type': 'EMAIL', 'pattern': re.compile(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b')},
            #Phone Numbers
            {'type': 'PHONE', 'pattern': re.compile(r'\+?\s*(?:\d{1,3}[\s-]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}\b')},
            #SSNs
            {'type': 'SSN', 'pattern': re.compile(r'\b\d{3}-\d{2}-\d{4}\b')},
            #Credit Card Numbers
            {'type': 'CREDIT_CARD', 'pattern': re.compile(r'\b(?:\d[ -]*?){13,16}\b')},
            #Dates of Birth
            {'type': 'DATE_OF_BIRTH', 'pattern': re.compile(r'\b(?:0[1-9]|[12]\d|3[01])[-/.](?:0[1-9]|1[0-2])[-/.](?:19|20)\d{2}\b')},
            #IP Addresses
            {'type': 'IP_ADDRESS', 'pattern': re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')},
            #PAN or official IDs
            {'type': 'PAN', 'pattern': re.compile(r'\b[A-Z]{5}\d{4}[A-Z]{1}\b')},
            #Aadhaar Numbers
            {'type': 'AADHAAR', 'pattern': re.compile(r'\b\d{4}\s?\d{4}\s?\d{4}\b')},
            #Physical or Mailing Addresses
            {'type': 'ADDRESS', 'pattern': re.compile(r'\b(?:Gat|Plot|Tower|Building|Flat|Door|Survey)\s+No\.?\s*[\w/\s,-]+', re.IGNORECASE)},
            #Concatenated / CamelCase Names (e.g., KushalSubbayyaHegde)
            {'type': 'PERSON', 'pattern': re.compile(r'\b[A-Z][a-z]+(?:[A-Z][a-z]+){2,}\b')}
        ]



    def get_fake_value(self, originalVal, piiType):
        key = originalVal.strip().upper()
        if key in self.pii_mapping:
            return self.pii_mapping[key]
            
        if piiType == 'EMAIL': 
            replacement = self.fake.email()
        elif piiType == 'PHONE': 
            replacement = self.fake.phone_number()
        elif piiType == 'ADDRESS': 
            replacement = self.fake.street_address() + ", " + self.fake.city()
        elif piiType in ['PAN', 'AADHAAR', 'SSN', 'CREDIT_CARD']: 
            replacement = "[REDACTED_ID]"
        elif piiType == 'DATE_OF_BIRTH':
            replacement = self.fake.date(pattern="%d/%m/%Y")
        elif piiType == 'IP_ADDRESS':
            replacement = self.fake.ipv4()
        elif piiType == 'PERSON':
            replacement = self.fake.name().upper() if originalVal.isupper() else self.fake.name()
        elif piiType == 'ORG':
            replacement = self.fake.company()
        else: 
            replacement = f"[REDACTED_{piiType}]"
            
        self.pii_mapping[key] = replacement
        return replacement

    #this function scans the text for PII and adds discovered names and organizations to the respective sets
    def scan_and_discover(self, text):
        
        if not text or len(text.strip()) < 4:
            return
            

        contact_match = re.search(r'(?:Contact\s*Person[:]?)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', text)
        if contact_match:
            name_str = contact_match.group(1).strip()
            if name_str.lower() not in self.exclude_terms:
                self.discovered_names.add(name_str)

        doc = self.nlp(text)
        for ent in doc.ents:
            entity_text = ent.text.strip()
            lower_val = entity_text.lower()

            
            if lower_val in self.exclude_terms or entity_text.isdigit() or len(entity_text) <= 3:
                continue
            if entity_text.isupper() and len(entity_text.split()) == 1:
                continue 
                
            if ent.label_ == "PERSON" and len(entity_text.split()) >= 2:
                self.discovered_names.add(entity_text)
            elif ent.label_ == "ORG" and len(entity_text.split()) >= 2:
                if not any(term in lower_val for term in ["sebi", "bse", "nse", "act", "supreme court"]):
                    self.discovered_orgs.add(entity_text)

    def redact_text(self, text):
        if not text:
            return text
        redacted = text
        
     
        for item in self.regex_patterns:
            redacted = item['pattern'].sub(lambda m: self.get_fake_value(m.group(0), item['type']), redacted)
            
        lower_text = redacted.lower()
        
       
        for org in self.sorted_orgs:
            if org.lower() in lower_text:
                escaped = re.escape(org)
                redacted = re.sub(rf'\b{escaped}\b', lambda m: self.get_fake_value(m.group(0), 'ORG'), redacted, flags=re.IGNORECASE)

     
        for name in self.sorted_names:
            if name.lower() in lower_text:
                escaped = re.escape(name)
                redacted = re.sub(rf'\b{escaped}\b', lambda m: self.get_fake_value(m.group(0), 'PERSON'), redacted, flags=re.IGNORECASE)
                
        return redacted

    def safe_replace_paragraph(self, paragraph, new_text):
        """Safely updates paragraph text while keeping run structure intact to prevent text merging."""
        if not paragraph.runs:
            paragraph.text = new_text
            return
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""

    def process_file(self, input_path, output_path):
        if not os.path.exists(input_path):
            print(f"Error: File not found at {input_path}")
            sys.exit(1)
            
        print("Pass 1: Automated Entity Discovery (Unseen Document Scan)...")
        doc = Document(input_path)
        
       
        for p in doc.paragraphs:
            self.scan_and_discover(p.text)
            
        
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    self.scan_and_discover(cell.text)
                    
        # print(f"Automatically identified {len(self.discovered_names)} names and {len(self.discovered_orgs)} companies.")
        self.sorted_names = sorted(list(self.discovered_names), key=len, reverse=True)
        self.sorted_orgs = sorted(list(self.discovered_orgs), key=len, reverse=True)
        
        print("Pass 2: Applying Safe Redaction Engine...")
        
        
        for p in doc.paragraphs:  #using this for loop to iterate through all the paragraphs in the document and then apply the redaction function to each paragraph. This allows for redaction of PII in the main body of the document.
            original = p.text
            if original.strip():
                redacted = self.redact_text(original)
                if original != redacted:
                    self.safe_replace_paragraph(p, redacted)
                        
        
        for t in doc.tables: #using this for loop to iterate through all the tables in the document and then iterate through each row and cell to access the paragraphs within each cell. This allows for redaction of PII in table cells as well.
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        original = p.text
                        if original.strip():
                            redacted = self.redact_text(original)
                            if original != redacted:
                                self.safe_replace_paragraph(p, redacted)

        doc.save(output_path)
        print(f"\nThe new document fully automated and clean file saved to {output_path}\n")